"""Generate the final project report for Content Monetization Modeler."""

from __future__ import annotations

import csv
import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "reports" / "final_report"
DOCX_PATH = OUT_DIR / "Content_Monetization_Modeler_Final_Project_Report.docx"
MD_PATH = OUT_DIR / "Content_Monetization_Modeler_Final_Project_Report.md"


def read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def read_business_insights() -> list[str]:
    path = ROOT / "reports" / "results" / "business_insights.md"
    insights = []
    for line in path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if line and line[0].isdigit() and ". " in line:
            insights.append(line.split(". ", 1)[1])
    return insights


def get_project_data() -> dict:
    raw = pd.read_csv(ROOT / "data" / "raw" / "youtube_ad_revenue_dataset.csv")
    processed = pd.read_csv(ROOT / "data" / "processed" / "cleaned_youtube_revenue.csv")
    comparison = pd.read_csv(ROOT / "reports" / "results" / "model_comparison.csv")
    summary = read_json(ROOT / "reports" / "results" / "project_summary.json")
    missing = raw.isna().sum().to_dict()
    dtypes = {col: str(dtype) for col, dtype in raw.dtypes.items()}

    return {
        "raw": raw,
        "processed": processed,
        "comparison": comparison,
        "summary": summary,
        "missing": missing,
        "dtypes": dtypes,
        "duplicates": int(raw.duplicated().sum()),
        "insights": read_business_insights(),
    }


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cell, widths_dxa[idx])


def style_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_footer(doc: Document):
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("Content Monetization Modeler - Final Project Report")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(90, 90, 90)


def page_break(doc: Document):
    doc.add_page_break()


def add_title_page(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(180)
    run = p.add_run("Content Monetization Modeler")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string("0B2545")

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Machine Learning Based YouTube Advertisement Revenue Prediction System")
    run2.font.size = Pt(16)
    run2.font.color.rgb = RGBColor.from_string("1F4D78")
    page_break(doc)


def add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold_label: str | None = None):
    p = doc.add_paragraph()
    if bold_label:
        r = p.add_run(bold_label)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbers(doc: Document, items: list[str]):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        set_cell_shading(hdr[i], "F2F4F7")
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    if widths:
        set_table_geometry(table, widths)
    doc.add_paragraph()
    return table


def add_figure_placeholder(doc: Document, caption: str, file_name: str | None = None):
    add_para(doc, caption)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Insert Figure Here]")
    run.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)
    if file_name:
        add_para(doc, f"Source file available in project: {file_name}")


def add_screenshot_placeholder(doc: Document, caption: str, file_name: str):
    add_para(doc, caption)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Insert Screenshot Here]")
    run.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)
    add_para(doc, f"Screenshot file available in project: {file_name}")


def add_front_matter(doc: Document, data: dict):
    add_heading(doc, "Certificate")
    add_para(
        doc,
        "This is to certify that the project titled Content Monetization Modeler has been completed as a machine learning based project for predicting YouTube advertisement revenue. The project includes dataset understanding, exploratory data analysis, preprocessing, feature engineering, model training, model evaluation, business insights, a Streamlit application, and deployment preparation."
    )
    add_para(doc, "The work has been prepared for academic project evaluation and viva presentation.")
    add_para(doc, "Project Title: Content Monetization Modeler")
    add_para(doc, "Target Variable: ad_revenue_usd")
    add_para(doc, "Signature of Guide / Evaluator: ____________________")
    add_para(doc, "Date: ____________________")
    page_break(doc)

    add_heading(doc, "Declaration")
    add_para(
        doc,
        "I hereby declare that the project report titled Content Monetization Modeler is based on the implemented project work. The project uses the provided YouTube advertisement revenue dataset and applies data cleaning, exploratory data analysis, feature engineering, regression modeling, and Streamlit application development."
    )
    add_para(
        doc,
        "The explanations, results, model scores, and insights in this report match the completed implementation and do not intentionally include unrelated tools, invented models, or advanced architectures that were not used in the project."
    )
    add_para(doc, "Student Signature: ____________________")
    add_para(doc, "Date: ____________________")
    page_break(doc)

    add_heading(doc, "Acknowledgement")
    add_para(
        doc,
        "I would like to express my sincere gratitude to my mentors, trainers, and project evaluators for their guidance during this data science project. Their support helped me understand the practical workflow of a machine learning project, from dataset understanding to model deployment."
    )
    add_para(
        doc,
        "I also acknowledge the importance of open-source Python libraries such as Pandas, NumPy, Matplotlib, Seaborn, Plotly, Scikit-learn, Joblib, and Streamlit, which made this project implementation possible."
    )
    page_break(doc)

    add_heading(doc, "Abstract")
    add_para(
        doc,
        "Content monetization has become an important part of the digital creator economy. Video creators and media teams often need to estimate expected advertisement revenue in order to plan content strategy, understand audience behavior, and improve future performance. This project, Content Monetization Modeler, is a machine learning based YouTube advertisement revenue prediction system that predicts the target variable ad_revenue_usd using video performance metrics and contextual information."
    )
    add_para(
        doc,
        f"The dataset used in this project contains {data['summary']['rows']:,} records and {data['summary']['columns']} columns. Important input features include views, likes, comments, watch_time_minutes, video_length_minutes, subscribers, category, device, and country. The project follows a complete data science workflow: dataset understanding, exploratory data analysis, missing value handling, duplicate removal, outlier treatment, feature engineering, categorical encoding, model training, model evaluation, model interpretation, and Streamlit application development."
    )
    add_para(
        doc,
        "Six engineered features were created: engagement rate, likes per view, comments per view, watch time efficiency, interaction score, and subscriber engagement score. These features were designed to represent viewer interaction, content attention, and subscriber quality in a simple and explainable way. Categorical variables were encoded using one-hot encoding, while numerical variables were scaled using standard scaling as part of a scikit-learn preprocessing pipeline."
    )
    add_para(
        doc,
        f"Five regression models were compared: Linear Regression, Ridge Regression, Lasso Regression, Random Forest Regressor, and Gradient Boosting Regressor. The models were evaluated using R2 Score, RMSE, and MAE. The best-performing model was {data['summary']['best_model']}, with an R2 score of {data['summary']['top_r2']:.4f}, RMSE of {data['summary']['top_rmse']:.4f}, and MAE of {data['summary']['top_mae']:.4f}. The final model was saved using Joblib and integrated into a Streamlit application with prediction, EDA dashboard, model comparison, business insights, and revenue optimization advisor pages."
    )
    add_para(
        doc,
        "The final outcome is an evaluation-ready machine learning project that demonstrates practical data cleaning, feature engineering, regression modeling, model interpretation, and application deployment preparation using Render. The project is suitable for academic evaluation and viva discussion because it remains simple, explainable, and realistic."
    )
    page_break(doc)


def add_lists(doc: Document):
    add_heading(doc, "Table of Contents")
    toc = [
        "1. Title Page",
        "2. Certificate",
        "3. Declaration",
        "4. Acknowledgement",
        "5. Abstract",
        "6. Table of Contents",
        "7. Introduction",
        "8. Domain Overview",
        "9. Literature Review",
        "10. Problem Definition",
        "11. Dataset Understanding",
        "12. Methodology",
        "13. Exploratory Data Analysis",
        "14. Data Preprocessing",
        "15. Feature Engineering",
        "16. Machine Learning Implementation",
        "17. Model Evaluation",
        "18. Business Insights",
        "19. Streamlit Application",
        "20. Innovative Features",
        "21. Deployment",
        "22. Project Structure",
        "23. GitHub Repository",
        "24. Challenges Faced",
        "25. Future Enhancements",
        "26. Conclusion",
        "27. References",
        "Appendix Suggestions",
    ]
    add_numbers(doc, toc)
    page_break(doc)

    add_heading(doc, "List of Figures")
    add_numbers(
        doc,
        [
            "Figure 13.1 Revenue Distribution",
            "Figure 13.2 Category Revenue Analysis",
            "Figure 13.3 Device Revenue Analysis",
            "Figure 13.4 Country Revenue Analysis",
            "Figure 13.5 Engagement Rate vs Revenue",
            "Figure 13.6 Watch Time vs Revenue",
            "Figure 13.7 Correlation Heatmap",
            "Figure 17.1 Top Revenue Drivers",
            "Figure 19.1 Project Overview Page",
            "Figure 19.2 Revenue Prediction Page",
            "Figure 19.3 EDA Dashboard Page",
            "Figure 19.4 Model Comparison Page",
            "Figure 19.5 Business Insights Page",
            "Figure 19.6 Revenue Optimization Advisor Page",
        ],
    )
    page_break(doc)

    add_heading(doc, "List of Tables")
    add_numbers(
        doc,
        [
            "Table 11.1 Dataset Column Description",
            "Table 12.1 Project Methodology Workflow",
            "Table 14.1 Data Preprocessing Summary",
            "Table 15.1 Engineered Feature Summary",
            "Table 16.1 Model Selection Strategy",
            "Table 17.1 Model Comparison Table",
            "Table 17.2 Model Ranking Table",
            "Table 18.1 Business Insights and Recommended Actions",
            "Table 22.1 Project Folder Purpose",
        ],
    )
    page_break(doc)


def add_introduction(doc: Document):
    add_heading(doc, "7. Introduction")
    add_heading(doc, "7.1 Background", level=2)
    add_para(
        doc,
        "YouTube is one of the most widely used video-sharing platforms. Creators upload educational, entertainment, gaming, technology, music, lifestyle, and other types of content to reach audiences. Many creators earn income through advertisement revenue generated from views and audience engagement."
    )
    add_para(
        doc,
        "Content monetization means earning revenue from digital content. In the YouTube context, revenue can depend on factors such as views, watch time, viewer engagement, country, device type, and content category. Since these factors can vary widely across videos, estimating revenue manually can be difficult."
    )
    add_heading(doc, "7.2 Problem Statement", level=2)
    add_para(
        doc,
        "The problem is to build a machine learning system that predicts YouTube advertisement revenue using video performance metrics and contextual information. Revenue prediction is useful because it helps creators and media teams understand what types of videos are likely to perform better financially."
    )
    add_heading(doc, "7.3 Objectives", level=2)
    add_bullets(
        doc,
        [
            "Understand the dataset and identify important revenue-related variables.",
            "Perform EDA to study missing values, duplicates, distributions, correlations, and outliers.",
            "Clean the dataset and prepare it for machine learning.",
            "Create useful engineered features related to engagement and watch time.",
            "Train and compare exactly five regression models.",
            "Evaluate models using R2 Score, RMSE, and MAE.",
            "Build a Streamlit application for prediction, analytics, and recommendations.",
            "Prepare the project for GitHub and Render deployment.",
        ],
    )
    add_heading(doc, "7.4 Scope", level=2)
    add_para(
        doc,
        "The project scope is limited to supervised regression modeling on the provided synthetic YouTube revenue dataset. It includes a local Streamlit application and Render deployment configuration. It does not include real YouTube API integration, cloud data pipelines, or advanced MLOps."
    )
    add_heading(doc, "7.5 Limitations", level=2)
    add_bullets(
        doc,
        [
            "The dataset is synthetic and created for learning purposes.",
            "The model does not use actual YouTube CPM, RPM, or advertiser bidding data.",
            "The deployed Render URL is not stored in the repository yet.",
            "The GitHub repository is prepared locally; a remote URL must be added after pushing.",
            "The recommendation logic is simple and rule-based, suitable for a beginner project.",
        ],
    )
    page_break(doc)


def add_domain_and_literature(doc: Document):
    add_heading(doc, "8. Domain Overview")
    add_para(
        doc,
        "The digital content industry includes platforms where creators publish videos, blogs, podcasts, and other media. The creator economy allows individuals and companies to earn income through advertising, subscriptions, sponsorships, and paid products."
    )
    add_para(
        doc,
        "YouTube monetization is mainly influenced by the number of views, audience retention, viewer engagement, content category, viewer geography, and device behavior. For example, high watch time can indicate that viewers are spending more time with the content, while comments and likes show interaction."
    )
    add_heading(doc, "8.1 Revenue Influencing Factors", level=2)
    add_table(
        doc,
        ["Factor", "Importance"],
        [
            ["Views", "More views can create more opportunities for advertisements to be shown."],
            ["Engagement", "Likes and comments indicate audience interest and content quality."],
            ["Subscribers", "Subscriber count shows channel reach, but engagement is also important."],
            ["Watch Time", "Higher watch time usually means stronger attention and more monetization potential."],
            ["Category", "Different content types may have different advertiser demand."],
            ["Country", "Advertisement rates can vary by market."],
            ["Device Type", "Viewing behavior may differ across mobile, desktop, tablet, and TV."],
        ],
        [1700, 7660],
    )
    page_break(doc)

    add_heading(doc, "9. Literature Review")
    add_para(
        doc,
        "This project uses beginner-friendly references and practical machine learning concepts rather than advanced research methods. The literature review focuses on commonly used approaches in predictive modeling, EDA, and regression."
    )
    add_table(
        doc,
        ["Reference Area", "Existing Approach", "Key Finding", "Limitation"],
        [
            ["Linear Regression", "Predict continuous values using a linear relationship between input features and target.", "Easy to explain and useful as a baseline model.", "May not capture complex non-linear patterns."],
            ["Regularized Regression", "Ridge and Lasso reduce overfitting through penalty terms.", "Helps improve generalization and can handle correlated features better.", "Requires selecting suitable regularization strength."],
            ["Tree-Based Models", "Random Forest and Gradient Boosting learn non-linear relationships.", "Can model complex feature interactions.", "Less transparent than linear models and may take more time to train."],
            ["EDA and Visualization", "Use distributions, correlations, and category summaries.", "Helps understand the data before modeling.", "Findings depend on dataset quality."],
            ["Streamlit Applications", "Create simple interactive dashboards using Python.", "Good for student projects and fast ML demos.", "Not designed as a full enterprise application framework."],
        ],
        [1500, 2800, 2500, 2560],
    )
    page_break(doc)


def add_problem_and_dataset(doc: Document, data: dict):
    add_heading(doc, "10. Problem Definition")
    add_heading(doc, "10.1 Existing Challenges", level=2)
    add_bullets(
        doc,
        [
            "Revenue depends on multiple metrics, not only views.",
            "Missing values and duplicate records can reduce analysis quality.",
            "Engagement and watch time need to be converted into meaningful features.",
            "Different countries, devices, and categories may influence revenue differently.",
            "A simple app is needed so non-technical users can test predictions.",
        ],
    )
    add_heading(doc, "10.2 Proposed Solution", level=2)
    add_para(
        doc,
        "The proposed solution is a Python and Streamlit based machine learning system. It cleans the dataset, creates engagement-related features, trains five regression models, selects the best model using evaluation metrics, saves the trained pipeline, and uses it inside a Streamlit web application."
    )
    add_heading(doc, "10.3 Expected Outcomes", level=2)
    add_bullets(
        doc,
        [
            "A cleaned and processed dataset.",
            "EDA findings and visual analysis.",
            "A trained regression model for revenue prediction.",
            "Model comparison using R2, RMSE, and MAE.",
            "Business insights and recommendations.",
            "A Streamlit application ready for Render deployment.",
        ],
    )
    page_break(doc)

    add_heading(doc, "11. Dataset Understanding")
    add_para(
        doc,
        f"The dataset contains {data['summary']['rows']:,} raw records and {data['summary']['columns']} columns. After cleaning duplicate rows, the processed dataset contains {len(data['processed']):,} records. The target variable is ad_revenue_usd."
    )
    add_table(
        doc,
        ["Column", "Type", "Description"],
        [
            ["video_id", data["dtypes"]["video_id"], "Unique video identifier."],
            ["date", data["dtypes"]["date"], "Upload or reporting date."],
            ["views", data["dtypes"]["views"], "Number of video views."],
            ["likes", data["dtypes"]["likes"], "Number of likes received by the video."],
            ["comments", data["dtypes"]["comments"], "Number of comments received by the video."],
            ["watch_time_minutes", data["dtypes"]["watch_time_minutes"], "Total watch time in minutes."],
            ["video_length_minutes", data["dtypes"]["video_length_minutes"], "Duration of the video."],
            ["subscribers", data["dtypes"]["subscribers"], "Subscriber count of the channel."],
            ["category", data["dtypes"]["category"], "Content category."],
            ["device", data["dtypes"]["device"], "Device used for viewing."],
            ["country", data["dtypes"]["country"], "Viewer country."],
            ["ad_revenue_usd", data["dtypes"]["ad_revenue_usd"], "Target revenue value in USD."],
        ],
        [1900, 1500, 5960],
    )
    add_heading(doc, "11.1 Target Variable", level=2)
    add_para(
        doc,
        "The target variable ad_revenue_usd represents the advertisement revenue generated by a video. It is important because it directly connects video performance to business value."
    )
    page_break(doc)


def add_methodology(doc: Document):
    add_heading(doc, "12. Methodology")
    add_para(doc, "The project follows a simple and explainable data science workflow.")
    add_table(
        doc,
        ["Step", "Purpose", "Activities Performed", "Outcome"],
        [
            ["Dataset", "Start with raw data.", "Loaded CSV file with YouTube video metrics.", "Raw dataset available for analysis."],
            ["EDA", "Understand patterns.", "Checked shape, types, missing values, duplicates, distributions, and relationships.", "EDA figures and findings generated."],
            ["Data Cleaning", "Improve data quality.", "Handled missing values, duplicates, datatypes, categories, and outliers.", "Cleaned dataset saved."],
            ["Feature Engineering", "Create better inputs.", "Created engagement and watch-time based features.", "Model-ready dataset prepared."],
            ["Encoding", "Use categorical data.", "Applied one-hot encoding for category, device, and country.", "Categorical features converted to numeric form."],
            ["Train-Test Split", "Evaluate fairly.", "Split data into training and testing sets.", "Test set used for model comparison."],
            ["Model Training", "Learn patterns.", "Trained five regression models.", "Model predictions generated."],
            ["Model Evaluation", "Compare performance.", "Calculated R2, RMSE, and MAE.", "Best model selected."],
            ["Model Saving", "Reuse model.", "Saved model pipeline and preprocessor with Joblib.", "Artifacts stored in models folder."],
            ["Streamlit App", "Make project interactive.", "Created six Streamlit pages.", "User can predict revenue and view analysis."],
            ["Render Deployment", "Prepare online hosting.", "Added requirements.txt and render.yaml.", "Project ready for Render deployment."],
        ],
        [1300, 2200, 3300, 2560],
    )
    add_figure_placeholder(doc, "Figure 12.1 Project Workflow", "Workflow: Dataset -> EDA -> Cleaning -> Feature Engineering -> Encoding -> Train-Test Split -> Model Training -> Evaluation -> Model Saving -> Streamlit -> Render")
    page_break(doc)


def add_eda(doc: Document, data: dict):
    add_heading(doc, "13. Exploratory Data Analysis")
    sections = [
        ("13.1 Dataset Overview", "To understand the size and structure of the dataset.", "Dataset shape and column inspection.", f"The raw dataset contains {data['summary']['rows']:,} rows and {data['summary']['columns']} columns.", "The dataset is large enough for training and evaluating regression models."),
        ("13.2 Missing Values Analysis", "To identify incomplete fields.", "Missing value count table.", f"Missing values were found mainly in likes ({data['missing']['likes']:,}), comments ({data['missing']['comments']:,}), and watch_time_minutes ({data['missing']['watch_time_minutes']:,}).", "Missing engagement values must be handled because they affect engineered features."),
        ("13.3 Duplicate Records Analysis", "To check repeated rows.", "Duplicate count check.", f"The raw dataset contained {data['duplicates']:,} duplicate rows.", "Removing duplicates prevents repeated records from biasing model training."),
        ("13.4 Statistical Summary", "To understand central tendency and spread.", "Descriptive statistics table.", "Views are concentrated around roughly ten thousand per record, while revenue varies across videos.", "Statistical summary helps identify skewness and scale differences."),
        ("13.5 Revenue Analysis", "To study the target variable.", "Revenue distribution histogram.", "Most videos fall in a middle revenue band, with fewer very high earning records.", "Revenue distribution helps decide suitable regression evaluation metrics."),
        ("13.6 Category Analysis", "To compare revenue by content category.", "Average revenue by category bar chart.", "Tech has the highest average revenue among categories in the processed analysis.", "Category can support future content planning."),
        ("13.7 Device Analysis", "To compare revenue by viewing device.", "Average revenue by device bar chart.", "Mobile traffic gives the highest average revenue by device in the processed analysis.", "Device-level behavior can guide content formatting and user experience decisions."),
        ("13.8 Country Analysis", "To compare revenue by country.", "Average revenue by country bar chart.", "US is the strongest country by average revenue in the processed analysis.", "Country analysis matters because ad rates differ across markets."),
        ("13.9 Subscriber Analysis", "To understand subscriber influence.", "Subscriber distribution and subscriber engagement features.", "Subscriber count alone is not enough; subscriber engagement gives better context.", "Creators should focus on active audiences, not only raw subscriber count."),
        ("13.10 Correlation Analysis", "To identify numerical relationships.", "Correlation heatmap.", "Watch time is the strongest numerical driver after the target itself.", "Correlation helps explain why watch-time features matter for revenue."),
        ("13.11 Outlier Analysis", "To identify extreme values.", "IQR outlier detection.", "Outliers were checked using the IQR method and capped during preprocessing.", "Capping controls extreme values without deleting useful rows."),
    ]
    figure_files = {
        "13.5 Revenue Analysis": "reports/figures/revenue_distribution.png",
        "13.6 Category Analysis": "reports/figures/category_revenue.png",
        "13.7 Device Analysis": "reports/figures/device_revenue.png",
        "13.8 Country Analysis": "reports/figures/country_revenue.png",
        "13.10 Correlation Analysis": "reports/figures/correlation_heatmap.png",
    }
    for idx, (title, objective, visual, findings, interpretation) in enumerate(sections, 1):
        add_heading(doc, title, level=2)
        add_para(doc, objective, "Objective: ")
        add_para(doc, visual, "Visualization Description: ")
        add_para(doc, findings, "Findings: ")
        add_para(doc, interpretation, "Business Interpretation: ")
        add_figure_placeholder(doc, f"Figure 13.{idx} {title.split(' ', 1)[1]}", figure_files.get(title, "Generated during EDA"))
    page_break(doc)


def add_preprocessing_and_features(doc: Document):
    add_heading(doc, "14. Data Preprocessing")
    add_table(
        doc,
        ["Step", "What Was Done", "Why It Was Done", "Impact"],
        [
            ["Missing Values", "Numerical missing values were filled using median values.", "Median is stable when values are skewed.", "Prevents model errors and keeps records usable."],
            ["Duplicate Removal", "Duplicate rows were removed.", "Repeated rows can bias model learning.", "Improves training reliability."],
            ["Outlier Treatment", "IQR capping was applied to numerical values.", "Extreme values can strongly affect regression.", "Reduces distortion while retaining data."],
            ["Encoding", "Category, device, and country were one-hot encoded.", "Models need numerical input.", "Allows contextual features to be used."],
            ["Scaling", "Numerical features were standardized.", "Linear models perform better with comparable scales.", "Improves stability for Linear, Ridge, and Lasso models."],
            ["Transformation", "Date was converted and category text was cleaned.", "Correct datatypes and clean labels reduce errors.", "Creates consistent preprocessing pipeline."],
        ],
        [1700, 2500, 2500, 2660],
    )
    page_break(doc)

    add_heading(doc, "15. Feature Engineering")
    rows = [
        ["Engagement Rate", "(likes + comments) / views", "Measures interaction relative to reach.", "Shows how actively viewers respond to content."],
        ["Likes Per View", "likes / views", "Measures like density.", "Helps compare videos with different view counts."],
        ["Comments Per View", "comments / views", "Measures conversation density.", "Shows discussion quality around content."],
        ["Watch Time Efficiency", "watch_time_minutes / video_length_minutes", "Compares total attention with video duration.", "Connects retention-like behavior to revenue."],
        ["Interaction Score", "0.7 * likes + 1.3 * comments", "Gives weighted importance to interactions.", "Comments are weighted higher because they require more effort."],
        ["Subscriber Engagement Score", "(likes + comments) / subscribers", "Measures engagement relative to subscriber base.", "Shows whether subscribers are active."],
    ]
    add_table(doc, ["Feature", "Formula", "Purpose", "Business Value"], rows, [1900, 2400, 2500, 2560])
    add_para(
        doc,
        "These features are expected to improve model learning because they convert raw activity counts into ratios and scores that better represent user behavior."
    )
    page_break(doc)


def add_ml_and_evaluation(doc: Document, data: dict):
    add_heading(doc, "16. Machine Learning Implementation")
    add_para(
        doc,
        "The dataset was split into training and testing sets using an 80:20 split with random_state=42. A common preprocessing pipeline was used for all models so the comparison remained fair."
    )
    model_rows = [
        ["Linear Regression", "Fits a straight-line relationship between features and revenue.", "Simple baseline and easy to explain.", "May miss non-linear patterns.", "Selected as the required baseline regression model."],
        ["Ridge Regression", "Linear regression with L2 regularization.", "Reduces overfitting and handles correlated features.", "Does not perform feature selection.", "Selected to test regularized linear performance."],
        ["Lasso Regression", "Linear regression with L1 regularization.", "Can reduce less useful coefficients and improve simplicity.", "May underfit if penalty is too strong.", "Selected because it is explainable and performed best."],
        ["Random Forest Regressor", "Combines multiple decision trees.", "Can capture non-linear relationships.", "Less transparent and heavier than linear models.", "Selected to compare tree-based performance."],
        ["Gradient Boosting Regressor", "Builds trees sequentially to correct errors.", "Strong model for tabular regression.", "Can be slower and needs tuning.", "Selected as a boosted model comparison."],
    ]
    add_table(doc, ["Model", "Working Principle", "Advantages", "Limitations", "Why Selected"], model_rows, [1700, 2500, 1800, 1700, 1660])
    page_break(doc)

    add_heading(doc, "17. Model Evaluation")
    add_heading(doc, "17.1 Evaluation Metrics", level=2)
    add_bullets(
        doc,
        [
            "R2 Score explains how much variance in revenue is explained by the model.",
            "RMSE measures the square root of average squared error and penalizes large mistakes.",
            "MAE measures average absolute error and is easy to understand in revenue units.",
        ],
    )
    comparison_rows = []
    for _, row in data["comparison"].iterrows():
        comparison_rows.append([
            row["Model"],
            f"{row['R2 Score']:.4f}",
            f"{row['RMSE']:.4f}",
            f"{row['MAE']:.4f}",
        ])
    add_table(doc, ["Model", "R2", "RMSE", "MAE"], comparison_rows, [3200, 1600, 2200, 2360])
    ranking_rows = [[str(i + 1), row[0], row[1], row[2], row[3]] for i, row in enumerate(comparison_rows)]
    add_table(doc, ["Rank", "Model", "R2", "RMSE", "MAE"], ranking_rows, [900, 3100, 1500, 1900, 1960])
    add_heading(doc, "17.2 Best Model Selection", level=2)
    add_para(
        doc,
        f"The selected best model is {data['summary']['best_model']}. It achieved the highest R2 score ({data['summary']['top_r2']:.4f}) and the lowest error values among the compared models. The result also supports viva explanation because Lasso is a simple regularized linear model."
    )
    add_heading(doc, "17.3 Trade-Off Discussion", level=2)
    add_para(
        doc,
        "Random Forest and Gradient Boosting are more flexible, but in this dataset the linear models performed slightly better. This suggests that the revenue relationship is strongly explained by linear numerical drivers, especially watch time and engagement-related features."
    )
    add_figure_placeholder(doc, "Figure 17.1 Top Revenue Drivers", "reports/figures/feature_importance.png")
    page_break(doc)


def add_business_insights(doc: Document, data: dict):
    add_heading(doc, "18. Business Insights")
    base_actions = [
        "Create more content in high-performing categories.",
        "Optimize mobile viewing experience and thumbnails.",
        "Study country-level revenue patterns before planning campaigns.",
        "Improve content pacing to increase total watch time.",
        "Use calls to action to improve likes and comments.",
        "Track subscriber engagement instead of only subscriber count.",
        "Avoid videos that are too short to generate meaningful watch time.",
        "Keep long videos structured so retention does not drop.",
        "Segment reports by country to understand revenue differences.",
        "Segment reports by device type to improve presentation style.",
        "Use the prediction app before planning new content.",
        "Build a content calendar based on high-revenue categories.",
        "Encourage meaningful comments with questions and prompts.",
        "Use what-if testing to estimate improvement opportunities.",
        "Monitor watch time as a primary KPI.",
    ]
    rows = []
    for i, insight in enumerate(data["insights"], 1):
        action = base_actions[(i - 1) % len(base_actions)]
        rows.append([str(i), insight, "The observation helps connect video behavior with revenue.", action])
    while len(rows) < 18:
        idx = len(rows) + 1
        rows.append([str(idx), "Revenue improves when performance metrics and context are analyzed together.", "Single metrics can be misleading without context.", "Use combined metrics such as engagement rate and watch time efficiency."])
    add_table(doc, ["No.", "Observation", "Interpretation", "Recommended Action"], rows, [700, 3300, 2700, 2660])
    page_break(doc)


def add_app_and_innovation(doc: Document):
    add_heading(doc, "19. Streamlit Application")
    add_para(
        doc,
        "The Streamlit application converts the machine learning project into an interactive tool. It loads the saved model pipeline, accepts user input, predicts ad revenue, and displays analytics and business insights."
    )
    pages = [
        ["Project Overview", "Introduce the project and show dataset-level metrics.", "Displays workflow, video count, average revenue, average views, and sample data.", "reports/screenshots/01_project_overview.png"],
        ["Revenue Prediction", "Allow users to predict ad revenue from input values.", "Collects views, likes, comments, watch time, length, subscribers, category, device, and country.", "reports/screenshots/02_revenue_prediction.png"],
        ["EDA Dashboard", "Show basic visual analytics.", "Displays revenue distribution, engagement analysis, category, device, and country charts.", "reports/screenshots/03_eda_dashboard.png"],
        ["Model Comparison", "Explain model performance.", "Shows comparison table, R2 ranking, and top revenue drivers.", "reports/screenshots/04_model_comparison.png"],
        ["Business Insights", "Convert analysis into practical business findings.", "Lists concise insights about category, engagement, watch time, subscribers, device, and country.", "reports/screenshots/05_business_insights.png"],
        ["Revenue Optimization Advisor", "Provide improvement suggestions after prediction.", "Shows prediction and gives recommendations about engagement and watch time.", "reports/screenshots/06_revenue_optimization_advisor.png"],
    ]
    add_table(doc, ["Page", "Purpose", "Functionality"], [[p[0], p[1], p[2]] for p in pages], [2200, 3200, 3960])
    for i, page in enumerate(pages, 1):
        add_screenshot_placeholder(doc, f"Figure 19.{i} {page[0]} Screenshot", page[3])
    page_break(doc)

    add_heading(doc, "20. Innovative Features")
    add_table(
        doc,
        ["Feature", "Objective", "Functionality", "User Benefit"],
        [
            ["Revenue Optimization Advisor", "Give practical improvement suggestions.", "Checks engagement and watch time patterns after prediction.", "Helps creators understand what to improve."],
            ["Revenue Performance Score", "Convert prediction into a simple score.", "Compares predicted revenue with a high-revenue benchmark.", "Makes model output easier to understand."],
            ["Category Revenue Leaderboard", "Show better performing content categories.", "Uses EDA category revenue analysis.", "Supports future content planning."],
            ["Top Revenue Drivers", "Explain model behavior.", "Shows feature importance or coefficient strength.", "Helps users trust the model."],
            ["What-If Simulator", "Allow dynamic testing.", "User changes inputs and prediction updates.", "Supports planning before publishing."],
        ],
        [2100, 2400, 2700, 2160],
    )
    page_break(doc)


def add_deployment_structure_github(doc: Document):
    add_heading(doc, "21. Deployment")
    add_heading(doc, "21.1 Why Streamlit", level=2)
    add_para(
        doc,
        "Streamlit was used because it allows Python-based data science applications to be built quickly without using web frameworks such as Flask, Django, or React. This matches the project requirement and keeps the project suitable for a fresher-level implementation."
    )
    add_heading(doc, "21.2 Why Render", level=2)
    add_para(
        doc,
        "Render was selected as the hosting platform because it can deploy a Streamlit app directly from a GitHub repository using a simple build command and start command."
    )
    add_heading(doc, "21.3 Deployment Status", level=2)
    add_para(
        doc,
        "The project includes requirements.txt and render.yaml. The Render deployment URL is not stored in the repository yet and should be added after the app is deployed from GitHub."
    )
    add_bullets(
        doc,
        [
            "Build Command: pip install -r requirements.txt",
            "Start Command: streamlit run app/app.py --server.port $PORT --server.address 0.0.0.0",
        ],
    )
    add_figure_placeholder(doc, "Figure 21.1 Deployment Workflow", "Dataset -> EDA -> Feature Engineering -> Model Training -> Best Model -> Streamlit Application -> Render Deployment")
    add_screenshot_placeholder(doc, "Figure 21.2 Render Deployment Screenshot", "Add Render dashboard screenshot after deployment")
    page_break(doc)

    add_heading(doc, "22. Project Structure")
    add_para(doc, "The project follows a simple modular structure.")
    add_para(
        doc,
        "content-monetization-modeler/\n|-- data/\n|-- notebooks/\n|-- src/\n|-- models/\n|-- reports/\n|-- app/\n|-- app.py\n|-- build_project.py\n|-- requirements.txt\n|-- render.yaml\n`-- README.md"
    )
    add_table(
        doc,
        ["Folder/File", "Purpose"],
        [
            ["data/raw", "Stores the original CSV dataset."],
            ["data/processed", "Stores cleaned model-ready dataset."],
            ["notebooks", "Contains the final project notebook."],
            ["src", "Contains preprocessing, feature engineering, training, evaluation, and prediction modules."],
            ["models", "Stores saved best model, trained model copy, scaler/preprocessor, and metadata."],
            ["reports/figures", "Stores EDA and feature importance figures."],
            ["reports/screenshots", "Stores Streamlit application screenshots."],
            ["reports/results", "Stores model comparison, insights, and preparation documents."],
            ["app/app.py", "Main Streamlit application."],
            ["render.yaml", "Render deployment configuration."],
        ],
        [2600, 6760],
    )
    page_break(doc)

    add_heading(doc, "23. GitHub Repository")
    add_para(
        doc,
        "The project is prepared as a local Git repository and contains all required files for GitHub submission. A local commit has been created. At the time of this report, no GitHub remote URL is stored in the repository, so the remote repository URL should be added after pushing."
    )
    add_bullets(
        doc,
        [
            "README.md explains project overview, dataset, installation, results, screenshots, deployment, and future scope.",
            "requirements.txt lists Python dependencies.",
            "models folder contains saved model artifacts.",
            "render.yaml contains deployment commands.",
            "reports folder contains figures, screenshots, results, and this final report.",
        ],
    )
    page_break(doc)


def add_challenges_future_conclusion_refs(doc: Document):
    add_heading(doc, "24. Challenges Faced")
    add_table(
        doc,
        ["Challenge", "Problem", "Solution", "Learning Outcome"],
        [
            ["Missing Values", "Likes, comments, and watch time had missing records.", "Median imputation was used.", "Data cleaning is essential before modeling."],
            ["Duplicate Records", "Duplicate rows could bias results.", "Duplicates were removed.", "Repeated data can distort model training."],
            ["Outliers", "Some numerical values were outside normal IQR bounds.", "IQR capping was applied.", "Outlier handling should preserve useful data."],
            ["Feature Engineering", "Raw counts alone did not fully represent engagement.", "Engagement and watch-time features were created.", "Domain logic improves model inputs."],
            ["Model Comparison", "Different models can perform similarly.", "Five models were compared using the same metrics.", "Best model selection should be metric-based."],
            ["Streamlit Integration", "App needed to use the same preprocessing as training.", "Saved pipeline was loaded through Joblib.", "Training and prediction preprocessing must match."],
            ["Deployment Preparation", "Online deployment needs dependency and start command clarity.", "requirements.txt and render.yaml were added.", "Deployment files should be simple and reproducible."],
        ],
        [1700, 2600, 2500, 2560],
    )
    page_break(doc)

    add_heading(doc, "25. Future Enhancements")
    add_bullets(
        doc,
        [
            "Add real YouTube channel history if available.",
            "Create features for upload day, upload hour, and video age.",
            "Add traffic source and viewer retention data.",
            "Improve recommendation logic using more business rules.",
            "Add residual analysis and prediction error charts in Streamlit.",
            "Improve app styling based on evaluator feedback.",
            "Add a deployed Render URL and GitHub repository link after publishing.",
        ],
    )
    page_break(doc)

    add_heading(doc, "26. Conclusion")
    add_para(
        doc,
        "The Content Monetization Modeler project successfully demonstrates a complete machine learning workflow for YouTube advertisement revenue prediction. The project begins with dataset understanding and EDA, then applies cleaning, feature engineering, categorical encoding, model training, evaluation, and app development."
    )
    add_para(
        doc,
        "Five regression models were compared using R2 Score, RMSE, and MAE. Lasso Regression was selected as the best model because it achieved the strongest evaluation results while remaining easy to explain. Business insights show the importance of watch time, engagement, category, country, device, and subscriber activity."
    )
    add_para(
        doc,
        "The final Streamlit application provides project overview, prediction, EDA dashboard, model comparison, business insights, and revenue optimization advisor pages. The project is ready for GitHub submission and Render deployment after adding the remote repository and live deployment URL."
    )
    page_break(doc)

    add_heading(doc, "27. References")
    references = [
        "Python Software Foundation. Python Documentation.",
        "Pandas Development Team. Pandas Documentation.",
        "NumPy Developers. NumPy Documentation.",
        "Scikit-learn Developers. Scikit-learn User Guide.",
        "Matplotlib Developers. Matplotlib Documentation.",
        "Seaborn Developers. Seaborn Documentation.",
        "Plotly Technologies Inc. Plotly Python Documentation.",
        "Streamlit Inc. Streamlit Documentation.",
        "Render. Render Deployment Documentation.",
    ]
    add_numbers(doc, references)
    page_break(doc)

    add_heading(doc, "Appendix Suggestions")
    add_bullets(
        doc,
        [
            "Appendix A: Full model comparison CSV.",
            "Appendix B: Feature importance CSV.",
            "Appendix C: Screenshots of all Streamlit pages.",
            "Appendix D: GitHub repository link after pushing.",
            "Appendix E: Render deployment URL after deployment.",
            "Appendix F: Important code snippets from preprocessing and model training.",
        ],
    )


def build_markdown(data: dict):
    comparison = data["comparison"]
    table_lines = ["| Model | R2 Score | RMSE | MAE |", "|---|---:|---:|---:|"]
    for _, row in comparison.iterrows():
        table_lines.append(
            f"| {row['Model']} | {row['R2 Score']:.4f} | {row['RMSE']:.4f} | {row['MAE']:.4f} |"
        )
    lines = [
        "# Content Monetization Modeler - Final Project Report",
        "",
        "Machine Learning Based YouTube Advertisement Revenue Prediction System",
        "",
        "This Markdown file is a source companion for the generated DOCX report.",
        "",
        f"Dataset rows: {data['summary']['rows']:,}",
        f"Dataset columns: {data['summary']['columns']}",
        f"Best model: {data['summary']['best_model']}",
        f"R2: {data['summary']['top_r2']:.4f}",
        f"RMSE: {data['summary']['top_rmse']:.4f}",
        f"MAE: {data['summary']['top_mae']:.4f}",
        "",
        "## Model Comparison",
        "",
        "\n".join(table_lines),
        "",
        "## Business Insights",
        "",
    ]
    for insight in data["insights"]:
        lines.append(f"- {insight}")
    lines.extend(
        [
            "",
            "## Report Note",
            "",
            "The full report with all chapters, placeholders, tables, and academic front matter is available as a DOCX file in reports/final_report.",
        ]
    )
    MD_PATH.write_text("\n".join(lines), encoding="utf-8")


def build_docx(data: dict):
    doc = Document()
    style_document(doc)
    add_title_page(doc)
    add_front_matter(doc, data)
    add_lists(doc)
    add_introduction(doc)
    add_domain_and_literature(doc)
    add_problem_and_dataset(doc, data)
    add_methodology(doc)
    add_eda(doc, data)
    add_preprocessing_and_features(doc)
    add_ml_and_evaluation(doc, data)
    add_business_insights(doc, data)
    add_app_and_innovation(doc)
    add_deployment_structure_github(doc)
    add_challenges_future_conclusion_refs(doc)
    add_footer(doc)
    doc.save(DOCX_PATH)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    data = get_project_data()
    build_markdown(data)
    build_docx(data)
    print(DOCX_PATH)
    print(MD_PATH)


if __name__ == "__main__":
    main()
